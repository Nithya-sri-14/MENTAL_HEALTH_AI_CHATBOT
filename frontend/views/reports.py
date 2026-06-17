from __future__ import annotations

import streamlit as st
import pandas as pd
from io import BytesIO
from api_client import get_patients, get_sessions, save_report, get_reports

try:
    from docx import Document
except Exception:
    Document = None


def export_docx(report_text: str) -> bytes:
    if Document is None:
        return report_text.encode("utf-8")
    doc = Document()
    for line in report_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_reports() -> None:
    st.subheader("Automated Clinical Reports")
    st.caption("Generate, edit, and export clinical summary sheets for patients based on screenings.")

    role = st.session_state.get("user_role")
    if role not in {"Admin", "Psychologist"}:
        st.info("Clinical report generation and reviews are restricted to psychologist and administrator accounts.")
        return

    # 1. Select Patient
    patients = get_patients()
    patient_names = [p["name"] for p in patients]
    
    patient_options = ["(none)"] + patient_names
    selected = st.session_state.get("active_patient") or "(none)"
    if selected not in patient_options:
        selected = "(none)"
        
    chosen = st.selectbox("Select Patient to Compile Report", patient_options, index=patient_options.index(selected))
    if chosen == "(none)":
        st.warning("Please select a patient in the dropdown or sidebar to view or compile reports.")
        return
    active_patient = chosen
    st.session_state.active_patient = chosen

    # Fetch patient profile details
    patient = next((p for p in patients if p["name"] == active_patient), None)
    if not patient:
        st.error("Patient profile not found.")
        return

    st.markdown("---")

    # Fetch latest screening sessions
    sessions = get_sessions(active_patient)
    if not sessions:
        st.warning(f"No clinical screenings found for {active_patient}. Please run an assessment first.")
        return

    # Get latest session details
    latest_session = sessions[0]
    
    # 2. Compile Report Draft
    st.markdown("### Generate Clinical Report")
    
    # Default draft structure
    draft_report = (
        f"# Clinical Workflow Report - {patient['name']}\n\n"
        f"- **Patient ID:** {patient.get('patient_id', 'N/A')}\n"
        f"- **Age:** {patient.get('age', 'N/A')}\n"
        f"- **Primary Concern:** {patient.get('primary_concern', 'N/A')}\n"
        f"- **Preferred Language:** {patient.get('language', 'English')}\n"
        f"- **Screening Score:** {latest_session.get('score', 0)} ({latest_session.get('risk_level', 'Low')} Risk)\n\n"
        f"## Session Summary\n"
        f"{latest_session.get('summary', 'No summary generated.')}\n\n"
        f"## Treatment Recommendations\n"
        f"{latest_session.get('recommendations', '- Follow standard checkup.')}\n\n"
        f"## Safety Declaration & Sign-off\n"
        f"Case escalated to psychiatrist review: {'YES' if latest_session.get('risk_level') in {'High', 'Critical'} else 'NO'}\n"
    )

    # Let the user review and edit the draft report
    report_editor = st.text_area("Review/Edit Markdown Report", value=draft_report, height=350)
    
    # 3. Actions
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Save Report to History"):
            success = save_report(active_patient, report_editor, latest_session.get('risk_level', 'Low'))
            if success:
                st.success("Report saved to history!")
                st.rerun()
            else:
                st.error("Failed to save report. Verify backend connection.")
                
    with col2:
        st.download_button(
            "Download TXT Report",
            data=report_editor.encode("utf-8"),
            file_name=f"{active_patient}_report_{latest_session.get('date')}.txt",
            mime="text/plain"
        )
        
    with col3:
        docx_bytes = export_docx(report_editor)
        st.download_button(
            "Download DOCX Report",
            data=docx_bytes,
            file_name=f"{active_patient}_report_{latest_session.get('date')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # 4. History list
    st.write("")
    st.subheader("Saved Reports History")
    reports = get_reports(active_patient)
    if reports:
        df = pd.DataFrame(reports)
        display_df = df[["id", "risk_level", "created_at"]].copy()
        st.dataframe(display_df, use_container_width=True)
        
        # Display selected report contents
        report_ids = [str(r["id"]) for r in reports]
        chosen_id = st.selectbox("Select report to view content", report_ids)
        selected_report = next((r for r in reports if str(r["id"]) == chosen_id), None)
        if selected_report:
            st.markdown("---")
            st.markdown("#### Report Preview")
            st.markdown(selected_report["report_text"])
    else:
        st.info("No saved reports found in history.")
